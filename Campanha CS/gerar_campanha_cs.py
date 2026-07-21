# -*- coding: utf-8 -*-
"""
Gerador de documento Word — Ação de Marketing CS Jul–Dez 2026
Efcaz | Gabriel Vital
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Paleta Efcaz
BRAND       = RGBColor(0x0E, 0x8F, 0xA3)
BRAND_DARK  = RGBColor(0x0A, 0x74, 0x88)
BRAND_LIGHT = RGBColor(0xE4, 0xF4, 0xF7)
AMBER       = RGBColor(0xD9, 0x8C, 0x00)
AMBER_LIGHT = RGBColor(0xFF, 0xF7, 0xE0)
GREEN       = RGBColor(0x1A, 0x9E, 0x6A)
GREEN_LIGHT = RGBColor(0xE6, 0xF7, 0xF0)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
TEXT        = RGBColor(0x15, 0x28, 0x30)
TEXT_2      = RGBColor(0x45, 0x60, 0x68)
GRAY_LIGHT  = RGBColor(0xF4, 0xF9, 0xFA)


def set_cell_bg(cell, color):
    hex_color = str(color)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def no_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell,
                top={'val': 'none'}, bottom={'val': 'none'},
                left={'val': 'none'}, right={'val': 'none'})


def run(para, text, bold=False, italic=False, color=None, size=None):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = 'Calibri'
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return r


def spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before * 20))
    sp.set(qn('w:after'), str(after * 20))
    pPr.append(sp)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            if j < len(widths_cm):
                cell.width = Cm(widths_cm[j])


# ── Documento ─────────────────────────────────────────────────
doc = Document()

sec = doc.sections[0]
sec.top_margin    = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.5)
sec.page_width    = Cm(21.0)
sec.page_height   = Cm(29.7)

sty = doc.styles['Normal']
sty.font.name = 'Calibri'
sty.font.size = Pt(10)
sty.font.color.rgb = TEXT


# ─── CAPA ────────────────────────────────────────────────────
ct = doc.add_table(rows=1, cols=1)
ct.alignment = WD_TABLE_ALIGNMENT.CENTER
cc = ct.cell(0, 0)
set_cell_bg(cc, BRAND)
no_borders(ct)

p = cc.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(p, before=36, after=6)
run(p, "Ação de Marketing CS", bold=True, color=WHITE, size=28)

p2 = cc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(p2, before=0, after=6)
run(p2, "Julho a Dezembro de 2026", color=WHITE, size=16)

p3 = cc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(p3, before=4, after=36)
run(p3, "Gabriel Vital  |  Customer Success Specialist  |  Efcaz",
    color=RGBColor(0xB2, 0xE2, 0xEA), size=10)

doc.add_paragraph()


# ─── SEÇÃO 1 — CONTEXTO ──────────────────────────────────────
h = doc.add_paragraph()
spacing(h, before=16, after=6)
run(h, "1.  Contexto da Campanha", bold=True, size=14, color=BRAND)

ctx = doc.add_table(rows=1, cols=1)
no_borders(ctx)
c = ctx.cell(0, 0)
set_cell_bg(c, GRAY_LIGHT)
set_cell_borders(c, left={'val': 'single', 'sz': 18, 'color': '0E8FA3'})

cp = c.paragraphs[0]
spacing(cp, before=8, after=4)
run(cp, "Esta não é uma campanha pontual de marketing. ", bold=True, color=TEXT, size=10)
run(cp,
    "É uma ação contínua de CS — comunicação periódica que traduz novidades da "
    "plataforma em impacto real para o cliente. Gabriel é a ponte entre a Efcaz "
    "e o cliente: não a empresa vendendo, mas o CS ajudando a usar melhor.",
    color=TEXT_2, size=10)

cp2 = c.add_paragraph()
spacing(cp2, before=4, after=8)
run(cp2, "Ferramenta: ", bold=True, color=TEXT, size=10)
run(cp2, "RD Station  |  ", color=TEXT_2, size=10)
run(cp2, "Cadência: ", bold=True, color=TEXT, size=10)
run(cp2, "~2× por mês  |  ", color=TEXT_2, size=10)
run(cp2, "Origem: ", bold=True, color=TEXT, size=10)
run(cp2, "Reunião de 17/07/2026 (Gabriel, Renato, Ricardo, Alison)", color=TEXT_2, size=10)

doc.add_paragraph()


# ─── SEÇÃO 2 — OBJETIVOS ─────────────────────────────────────
h = doc.add_paragraph()
spacing(h, before=14, after=6)
run(h, "2.  Objetivos", bold=True, size=14, color=BRAND)

objs = [
    ("Engajamento da base",
     "Reativar contatos frios, medir abertura e atualizar base de e-mails.",
     BRAND, BRAND_LIGHT),
    ("Adoção de módulos",
     "Aproximar clientes de funcionalidades contratadas mas subutilizadas.",
     AMBER, AMBER_LIGHT),
    ("Expansão e renovação",
     "Abrir conversa de expansão antes do momento de renovação, sem pressão.",
     GREEN, GREEN_LIGHT),
]

ot = doc.add_table(rows=1, cols=3)
ot.alignment = WD_TABLE_ALIGNMENT.CENTER
no_borders(ot)
for i, (title, desc, color, bg) in enumerate(objs):
    cell = ot.cell(0, i)
    set_cell_bg(cell, bg)
    set_cell_borders(cell,
        top={'val': 'none'}, bottom={'val': 'none'},
        left={'val': 'none'}, right={'val': 'none'})
    pt = cell.paragraphs[0]
    spacing(pt, before=10, after=4)
    run(pt, title, bold=True, color=color, size=10)
    pd = cell.add_paragraph()
    spacing(pd, before=2, after=10)
    run(pd, desc, color=TEXT_2, size=9)
set_col_widths(ot, [5.3, 5.3, 5.3])
doc.add_paragraph()


# ─── SEÇÃO 3 — PREMISSAS ─────────────────────────────────────
h = doc.add_paragraph()
spacing(h, before=14, after=6)
run(h, "3.  Premissas da Campanha", bold=True, size=14, color=BRAND)

premissas = [
    ("Benefício primeiro",
     "Traduzir novidades em impacto real — não comunicar funcionalidade técnica, "
     "mas o que ela muda na operação do cliente."),
    ("Gabriel como ponte",
     "Comunicação parte do CS, não da empresa. Cria proximidade e reduz "
     "percepção de e-mail comercial."),
    ("Fase 1 geral, Fase 2 segmentada",
     "E-mails 1 e 2 para toda a base. A partir do e-mail 3, segmentação "
     "por perfil de uso da plataforma."),
    ("CTA de baixo atrito",
     "Agenda aberta do Gabriel como principal chamada para ação — fácil de "
     "clicar, sem comprometimento."),
    ("Volume controlado",
     "Calendário com slots abertos em periodicidade maior para evitar "
     "avalanche de agendamentos simultâneos."),
    ("Integração com CS",
     "Dados de abertura de e-mail viram gatilhos para WhatsApp, calls e "
     "identificação de oportunidades de expansão."),
]

for title, desc in premissas:
    p = doc.add_paragraph(style='List Bullet')
    spacing(p, before=3, after=3)
    run(p, title + ": ", bold=True, color=BRAND_DARK, size=10)
    run(p, desc, color=TEXT_2, size=10)

doc.add_paragraph()


# ─── SEÇÃO 4 — CRONOGRAMA ────────────────────────────────────
h = doc.add_paragraph()
spacing(h, before=14, after=8)
run(h, "4.  Cronograma Jul–Dez 2026", bold=True, size=14, color=BRAND)

headers = ["#", "Período", "Tema", "Hook (pergunta-âncora)", "Segmento", "Objetivo"]
rows_data = [
    ("1",  "Jul / agora",
     "Módulo de Contratos",
     '"Onde ficam os contratos dos seus fornecedores hoje?"',
     "Geral — toda a base",
     "Validar abertura, mapear contatos, medir engajamento"),

    ("2",  "Ago / 1ª quinz.",
     "Background Check / Compliance",
     '"Você sabe se seus fornecedores estão regulares hoje?"',
     "Geral (2ª rodada)",
     "Identificar interesse em compliance, mapear quem abre 2×"),

    ("3",  "Ago / 2ª quinz.",
     "Segmentação — 2 versões",
     "A: Terceiros  |  B: RFI",
     "A: só homologação  |  B: com Terceiros",
     "Primeira comunicação segmentada por perfil de uso"),

    ("4",  "Set / 1ª quinz.",
     "Avaliação de Performance (RFI)",
     '"Como você documenta quando um fornecedor entrega mal?"',
     "Clientes sem RFI ativo",
     "Abrir conversa de expansão de módulo"),

    ("5",  "Set / 2ª quinz.",
     "Webinar — Novidades da plataforma",
     "Convite para toda a base | gravação disponível",
     "Toda a base",
     "Reativar frios, gerar cases, insumo ao comercial"),

    ("6",  "Out / 1ª quinz.",
     "Dossiê + Score de Confiabilidade",
     '"Você já usou o dossiê para tomar uma decisão?"',
     "Clientes com baixo uso do dossiê",
     "Engajar recurso subutilizado, abrir conversa de valor"),

    ("7",  "Out / 2ª quinz.",
     "Gestão de Terceiros",
     '"Quantos colaboradores estão ativos e regularizados hoje?"',
     "Clientes com Terceiros ativo",
     "Reforçar adoção, identificar gaps de regularização"),

    ("8",  "Nov / 1ª quinz.",
     "Ocorrências + Planos de Ação",
     '"Quando algo dá errado com um fornecedor, como você registra?"',
     "Clientes sem Ocorrências",
     "Ampliar uso além da homologação"),

    ("9",  "Nov / 2ª quinz.",
     "Relatórios Gerenciais",
     '"Seu gestor consegue ver a saúde da base em 5 minutos?"',
     "Decisores / gestores",
     "Criar visibilidade com liderança, preparar renovações"),

    ("10", "Dezembro",
     "Balanço 2026 + Olhar 2027",
     '"O que sua empresa conquistou na gestão de fornecedores?"',
     "Toda a base",
     "Relacionamento, retenção, NPS, abertura de renovações"),
]

col_w = [0.7, 2.4, 3.0, 4.2, 3.5, 4.2]
tbl = doc.add_table(rows=1 + len(rows_data), cols=6)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for j, h_txt in enumerate(headers):
    cell = tbl.cell(0, j)
    set_cell_bg(cell, BRAND)
    set_cell_borders(cell,
        top={'val': 'none'}, bottom={'val': 'none'},
        left={'val': 'single', 'sz': 4, 'color': 'FFFFFF'},
        right={'val': 'single', 'sz': 4, 'color': 'FFFFFF'})
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    spacing(p, before=5, after=5)
    run(p, h_txt, bold=True, color=WHITE, size=9)

PHASE_BG = {
    "1":  (BRAND_LIGHT, BRAND),
    "2":  (BRAND_LIGHT, BRAND),
    "3":  (AMBER_LIGHT, AMBER),
    "4":  (AMBER_LIGHT, AMBER),
    "5":  (GREEN_LIGHT, GREEN),
    "6":  (AMBER_LIGHT, AMBER),
    "7":  (AMBER_LIGHT, AMBER),
    "8":  (AMBER_LIGHT, AMBER),
    "9":  (GREEN_LIGHT, GREEN),
    "10": (GREEN_LIGHT, GREEN),
}

white_border = {'val': 'single', 'sz': 4, 'color': 'FFFFFF'}

for i, row_data in enumerate(rows_data):
    num = row_data[0]
    bg, accent = PHASE_BG.get(num, (GRAY_LIGHT, TEXT_2))
    for j, val in enumerate(row_data):
        cell = tbl.cell(i + 1, j)
        set_cell_bg(cell, accent if j == 0 else bg)
        set_cell_borders(cell,
            top=white_border, bottom=white_border,
            left=white_border, right=white_border)
        p = cell.paragraphs[0]
        spacing(p, before=4, after=4)
        color = WHITE if j == 0 else (TEXT if j <= 2 else TEXT_2)
        run(p, val, bold=(j <= 1), color=color, size=9 if j > 0 else 10)

set_col_widths(tbl, col_w)
doc.add_paragraph()


# ─── SEÇÃO 5 — FASES ─────────────────────────────────────────
h = doc.add_paragraph()
spacing(h, before=14, after=6)
run(h, "5.  Fases de Segmentação", bold=True, size=14, color=BRAND)

fases = [
    ("F1", "Fase 1 — Geral (E-mails 1 e 2)",
     "E-mails 1 e 2 disparados para toda a base de clientes.\n"
     "Objetivo: validar contatos ativos, medir abertura e identificar bounces.\n"
     "Meta: taxa de abertura >25%  |  bounce <5%."),
    ("F2", "Fase 2 — Segmentado por uso (E-mail 3 em diante)",
     "Perfil A: cliente usa só homologação, sem Terceiros contratado\n"
     "Perfil B: cliente usa homologação + Terceiros\n"
     "Perfil C: cliente usa homologação + Terceiros + RFI\n"
     "Decisores sem acesso à plataforma: segmento exclusivo para visibilidade gerencial"),
    ("F3", "Fase 3 — Por comportamento (Nov/Dez)",
     "Abriu 3+ e-mails: cliente engajado — prioridade para conversa de expansão\n"
     "Abriu mas nunca agendou: acionar via WhatsApp ou ligação\n"
     "Nunca abriu: verificar contato, tentar outro canal (WhatsApp, call direta)"),
]

for label, title, desc in fases:
    ft = doc.add_table(rows=1, cols=2)
    no_borders(ft)
    set_col_widths(ft, [1.6, 15.9])
    bc = ft.cell(0, 0)
    set_cell_bg(bc, BRAND)
    bp = bc.paragraphs[0]
    spacing(bp, before=14, after=4)
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(bp, label, bold=True, color=WHITE, size=16)
    dc = ft.cell(0, 1)
    set_cell_bg(dc, BRAND_LIGHT)
    cp = dc.paragraphs[0]
    spacing(cp, before=8, after=2)
    run(cp, title, bold=True, color=BRAND_DARK, size=10)
    cp2 = dc.add_paragraph()
    spacing(cp2, before=3, after=8)
    run(cp2, desc, color=TEXT_2, size=9)
    doc.add_paragraph()


# ─── SEÇÃO 6 — E-MAIL 1 ──────────────────────────────────────
h = doc.add_paragraph()
spacing(h, before=14, after=6)
run(h, "6.  E-mail 1 — Módulo de Contratos", bold=True, size=14, color=BRAND)

mt = doc.add_table(rows=1, cols=4)
no_borders(mt)
set_col_widths(mt, [2.8, 4.2, 2.5, 8.0])
meta = [("Disparo:", "Geral — toda a base"), ("Previsto:", "Semana de 21/07/2026")]
for i, (lbl, val) in enumerate(meta):
    lc = mt.cell(0, i * 2)
    vc = mt.cell(0, i * 2 + 1)
    set_cell_bg(lc, BRAND)
    set_cell_bg(vc, BRAND_LIGHT)
    pl = lc.paragraphs[0]; spacing(pl, before=5, after=5)
    run(pl, lbl, bold=True, color=WHITE, size=9)
    pv = vc.paragraphs[0]; spacing(pv, before=5, after=5)
    run(pv, val, color=TEXT, size=9)

doc.add_paragraph()

sp = doc.add_paragraph()
spacing(sp, before=6, after=4)
run(sp, "Assunto: ", bold=True, color=BRAND_DARK, size=10)
run(sp, "Uma novidade que pode mudar como você controla contratos de fornecedores",
    color=TEXT, size=10)

et = doc.add_table(rows=1, cols=1)
no_borders(et)
ec = et.cell(0, 0)
set_cell_bg(ec, GRAY_LIGHT)
set_cell_borders(ec, left={'val': 'single', 'sz': 18, 'color': '0E8FA3'})

email_paras = [
    "Olá, [Nome],",
    "Meu nome é Gabriel, e sou o responsável pela sua conta aqui na Efcaz.",
    "A partir de agora, além do suporte que você já tem, vou te enviar comunicações "
    "periódicas com uma perspectiva diferente das news de produto: o que as novidades "
    "da plataforma mudam na prática da sua operação.",
]
for idx, text in enumerate(email_paras):
    p = ec.paragraphs[0] if idx == 0 else ec.add_paragraph()
    spacing(p, before=3, after=2)
    run(p, text, color=TEXT, size=10)

sep = ec.add_paragraph()
spacing(sep, before=8, after=2)
run(sep, "A estreia: Módulo de Contratos", bold=True, color=BRAND, size=10)

for line in [
    "Uma pergunta rápida: onde ficam os contratos dos seus fornecedores hoje?",
    "Para muitas empresas, a resposta é e-mail, pasta compartilhada ou planilha. "
    "Funciona — até o dia em que um contrato vence sem que ninguém perceba, ou quando "
    "você precisa encontrar uma vigência específica às pressas.",
    "Acabamos de liberar o Módulo de Contratos na plataforma. Com ele, você centraliza "
    "em um só lugar:",
]:
    p = ec.add_paragraph()
    spacing(p, before=4, after=2)
    run(p, line, color=TEXT, size=10)

for item in [
    "Vigências, aditivos e valores contratuais",
    "Documentos obrigatórios vinculados ao contrato",
    "Visibilidade antecipada de vencimentos",
]:
    p = ec.add_paragraph(style='List Bullet')
    spacing(p, before=2, after=2)
    run(p, item, color=TEXT_2, size=10)

p = ec.add_paragraph()
spacing(p, before=4, after=3)
run(p, "Tudo no mesmo ambiente onde você já gerencia seus fornecedores.", color=TEXT, size=10)

sep2 = ec.add_paragraph()
spacing(sep2, before=8, after=2)
run(sep2, "Isso faz sentido para a sua operação?", bold=True, color=BRAND, size=10)

for line in [
    "Cada empresa trabalha de um jeito — por isso, antes de qualquer coisa, queria "
    "entender se vale a pena te mostrar como isso ficaria configurado no ambiente de vocês.",
    "Se quiser, reserva 20 minutos aqui: [link da agenda]",
    "Ou responde esse e-mail que eu te retorno.",
]:
    p = ec.add_paragraph()
    spacing(p, before=4, after=3)
    run(p, line, color=TEXT, size=10)

cl = ec.add_paragraph()
spacing(cl, before=8, after=10)
run(cl, "Um abraço,\n", color=TEXT, size=10)
run(cl, "Gabriel Vital", bold=True, color=BRAND_DARK, size=10)
run(cl, "\nCustomer Success | Efcaz", color=TEXT_2, size=9)

doc.add_paragraph()


# ─── SEÇÃO 7 — PRÓXIMOS PASSOS ───────────────────────────────
h = doc.add_paragraph()
spacing(h, before=14, after=6)
run(h, "7.  Próximos Passos", bold=True, size=14, color=BRAND)

steps = [
    ("22/07 (quarta)",       "Compartilhar pasta com Renato, Ricardo e Alison"),
    ("23–24/07 (qui/sex)",   "Sessão com Alison para rodar E-mail 1 no RD Station + criar usuário do Gabriel"),
    ("Próxima reunião CX",   "Verificar integração CustomerX × RD Station"),
    ("A definir c/ Ricardo", "Confirmar data do webinar de setembro"),
]

pt = doc.add_table(rows=len(steps), cols=2)
no_borders(pt)
set_col_widths(pt, [4.5, 13.5])

for i, (date, action) in enumerate(steps):
    dc = pt.cell(i, 0)
    ac = pt.cell(i, 1)
    set_cell_bg(dc, BRAND if i % 2 == 0 else BRAND_DARK)
    set_cell_bg(ac, BRAND_LIGHT if i % 2 == 0 else WHITE)
    wb = {'val': 'single', 'sz': 4, 'color': 'FFFFFF'}
    lb = {'val': 'single', 'sz': 4, 'color': 'D0E8ED'}
    set_cell_borders(dc, top=wb, bottom=wb, left={'val': 'none'}, right={'val': 'none'})
    set_cell_borders(ac, top=lb, bottom=lb, left={'val': 'none'}, right={'val': 'none'})
    pd = dc.paragraphs[0]; spacing(pd, before=5, after=5)
    run(pd, date, bold=True, color=WHITE, size=9)
    pa = ac.paragraphs[0]; spacing(pa, before=5, after=5)
    run(pa, action, color=TEXT, size=10)

doc.add_paragraph()


# ─── RODAPÉ ───────────────────────────────────────────────────
hr = doc.add_table(rows=1, cols=1)
no_borders(hr)
hc = hr.cell(0, 0)
set_cell_bg(hc, BRAND)
spacing(hc.paragraphs[0], before=1, after=1)

ft = doc.add_paragraph()
ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacing(ft, before=6, after=0)
run(ft, "Efcaz  |  Ação de Marketing CS  |  Criado em 21/07/2026  |  Gabriel Vital",
    color=TEXT_2, size=8)


# ─── Salvar ───────────────────────────────────────────────────
output = r"c:\Users\gabriel.evangelista\Documents\ClaudeGL\Campanha CS\campanha_cs_21-07-2026.docx"
doc.save(output)
print(f"Documento gerado: {output}")
